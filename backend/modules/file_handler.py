"""
file_handler.py
Parses uploaded files (CSV, Excel, JSON, PDF, DOCX) into Pandas DataFrames.
"""

import io
import json
import pandas as pd
import PyPDF2
import docx
from werkzeug.datastructures import FileStorage


class FileHandler:

    SUPPORTED = {"csv", "xlsx", "xls", "json", "pdf", "docx"}

    @staticmethod
    def get_extension(filename: str) -> str:
        return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    @classmethod
    def parse(cls, file) -> dict:
        """
        Parse an uploaded file-like object (FileStorage or BytesIO with .filename).
        Returns a dict with dataframe, metadata, etc.
        """
        filename = getattr(file, 'filename', 'unknown')
        ext = cls.get_extension(filename)

        if ext not in cls.SUPPORTED:
            raise ValueError(f"Unsupported file type: .{ext}. Supported: {cls.SUPPORTED}")

        # Seek to start in case it's a BytesIO at EOF
        if hasattr(file, 'seek'):
            file.seek(0)
        raw_bytes = file.read()
        if not raw_bytes:
            raise ValueError("File is empty (0 bytes received).")

        return cls.parse_bytes(raw_bytes, filename)

    @classmethod
    def parse_bytes(cls, raw_bytes: bytes, filename: str) -> dict:
        """
        Parse raw bytes directly — preferred method for FastAPI UploadFile.
        Returns a dict with dataframe, metadata, etc.
        """
        if not raw_bytes:
            raise ValueError("File is empty (0 bytes received).")

        ext = cls.get_extension(filename)
        if ext not in cls.SUPPORTED:
            raise ValueError(f"Unsupported file type: .{ext}. Supported: {cls.SUPPORTED}")

        df = cls._parse_by_type(raw_bytes, ext, filename)

        return {
            "dataframe":    df,
            "file_type":    ext,
            "filename":     filename,
            "row_count":    len(df),
            "column_count": len(df.columns),
            "columns":      list(df.columns),
            "dtypes":       {col: str(dtype) for col, dtype in df.dtypes.items()},
            "raw_bytes":    raw_bytes,
        }

    @classmethod
    def _parse_by_type(cls, raw_bytes: bytes, ext: str, filename: str) -> pd.DataFrame:
        if ext == "csv":
            return cls._parse_csv(raw_bytes)
        elif ext in ("xlsx", "xls"):
            return cls._parse_excel(raw_bytes)
        elif ext == "json":
            return cls._parse_json(raw_bytes)
        elif ext == "pdf":
            return cls._parse_pdf(raw_bytes)
        elif ext == "docx":
            return cls._parse_docx(raw_bytes)
        raise ValueError(f"Parser not implemented for: {ext}")

    @staticmethod
    def _parse_csv(raw: bytes) -> pd.DataFrame:
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return pd.read_csv(io.BytesIO(raw), encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode CSV file.")

    @staticmethod
    def _parse_excel(raw: bytes) -> pd.DataFrame:
        return pd.read_excel(io.BytesIO(raw))

    @staticmethod
    def _parse_json(raw: bytes) -> pd.DataFrame:
        data = json.loads(raw.decode("utf-8"))
        if isinstance(data, list):
            return pd.DataFrame(data)
        elif isinstance(data, dict):
            # Try to normalize nested JSON
            return pd.json_normalize(data)
        raise ValueError("JSON must be a list of records or a nested object.")

    @staticmethod
    def _parse_pdf(raw: bytes) -> pd.DataFrame:
        """
        Extract text/tables from PDF. Tries multiple libraries with graceful fallbacks.
        """
        lines = []

        # ── Strategy 1: pypdf (modern, handles more formats) ── #
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            for page in reader.pages:
                text = page.extract_text() or ""
                lines.extend([l for l in text.split("\n") if l.strip()])
            if lines:
                rows = [line.split() for line in lines if line.strip()]
                if rows:
                    max_cols = max(len(r) for r in rows)
                    rows = [r + [""] * (max_cols - len(r)) for r in rows]
                    headers = [f"col_{i}" if not rows[0][i] else rows[0][i] for i in range(max_cols)]
                    return pd.DataFrame(rows[1:], columns=headers) if len(rows) > 1 else pd.DataFrame({"text": lines})
        except Exception:
            pass

        # ── Strategy 2: PyPDF2 (legacy) ── #
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            for page in reader.pages:
                text = page.extract_text() or ""
                lines.extend([l for l in text.split("\n") if l.strip()])
        except Exception:
            pass

        # ── Strategy 3: pdfplumber (table-aware) ── #
        if not lines:
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(raw)) as pdf:
                    for page in pdf.pages:
                        # Try tables first
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                lines.append(" | ".join([str(c or "") for c in row]))
                        # Fallback to text
                        if not tables:
                            text = page.extract_text() or ""
                            lines.extend([l for l in text.split("\n") if l.strip()])
            except Exception:
                pass

        if not lines:
            raise ValueError(
                "Could not extract content from this PDF. "
                "The file may be scanned, image-only, encrypted, or corrupted. "
                "Please try a CSV or Excel version of your data."
            )

        return pd.DataFrame({"text": lines})


    @staticmethod
    def _parse_docx(raw: bytes) -> pd.DataFrame:
        """Extract tables from DOCX; fall back to paragraph text."""
        doc = docx.Document(io.BytesIO(raw))
        # Try tables first
        if doc.tables:
            table = doc.tables[0]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            data = []
            for row in table.rows[1:]:
                data.append([cell.text.strip() for cell in row.cells])
            return pd.DataFrame(data, columns=headers)
        # Fall back to paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return pd.DataFrame({"text": paragraphs})


    @staticmethod
    def get_preview(df: pd.DataFrame, n: int = 10) -> list[dict]:
        """Return first n rows as a list of dicts (JSON-serializable)."""
        preview = df.head(n).copy()
        # Replace NaN with None for JSON compatibility
        preview = preview.where(pd.notna(preview), None)
        return preview.to_dict(orient="records")
